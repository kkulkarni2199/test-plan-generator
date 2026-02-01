"""
Endpoint 5 & 6: DVP Document Generation and Download
Generates Excel DVP matching reference format
"""
from fastapi import APIRouter, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
from typing import List, Dict, Any
import uuid
from datetime import datetime
from pathlib import Path
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models.api_models import (
    DVPGenerationRequest,
    DVPGenerationResponse,
    JobStatus,
    JobStatusResponse
)
from app.config import settings
from loguru import logger

router = APIRouter()

# In-memory job storage
dvp_jobs = {}
generated_dvps = {}  # Store DVP metadata

class DVPGenerator:
    """
    Generates DVP Excel document matching reference format
    """

    def __init__(self):
        self.workbook = None

    def generate_dvp(self, component_profile: Dict[str, Any],
                    test_cases: List[Dict[str, Any]],
                    include_traceability: bool = True) -> str:
        """
        Generate complete DVP Excel document
        """
        logger.info(f"Generating DVP for: {component_profile.get('name')}")

        # Create workbook
        self.workbook = Workbook()

        # Remove default sheet
        if 'Sheet' in self.workbook.sheetnames:
            del self.workbook['Sheet']

        # Sheet 1: Annex B - Test Matrix
        self._create_test_matrix_sheet(component_profile, test_cases)

        # Sheet 2: Test Sequence
        self._create_test_sequence_sheet(test_cases)

        # Sheet 3: Traceability Matrix (if requested)
        if include_traceability:
            self._create_traceability_sheet(test_cases)

        # Sheet 4: Source References
        self._create_references_sheet(test_cases)

        # Save file
        output_filename = f"DVP_{component_profile.get('name', 'Component').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        output_path = Path(settings.output_dir) / output_filename

        self.workbook.save(str(output_path))
        logger.info(f"DVP saved to: {output_path}")

        return str(output_path)

    def generate_dvp_docx(self, component_profile: Dict[str, Any],
                         test_cases: List[Dict[str, Any]],
                         include_traceability: bool = True) -> str:
        """
        Generate descriptive DVP Word document
        """
        logger.info(f"Generating DVP Docx for: {component_profile.get('name')}")
        
        document = Document()
        
        # Add a border to the page (simulated with a table or just a nice header)
        header = document.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = f"CONFIDENTIAL - {component_profile.get('name', 'DVP')}"
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Title Page
        document.add_paragraph("\n\n\n\n")
        title = document.add_heading(f"Design Verification Plan (DVP)", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        document.add_paragraph("\n")
        subtitle = document.add_paragraph(f"{component_profile.get('name', 'Component')}")
        subtitle.style = 'Title'
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        document.add_paragraph(f"Type: {component_profile.get('type', '')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(f"Application: {component_profile.get('application', '')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(f"Document Version: 1.0").alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add footer with page numbers
        footer = document.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"DVP-GEN System | Page "
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        document.add_page_break()
        
        # 1. Introduction
        document.add_heading('1. Introduction', level=1)
        document.add_paragraph(
            f"This document outlines the Design Verification Plan for the {component_profile.get('name')} {component_profile.get('type')}. "
            f"The purpose of this plan is to verify that the component meets all specified requirements for {component_profile.get('application')} application."
        )
        
        # 1.1 Component Details
        document.add_heading('1.1 Component Details', level=2)
        table = document.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        # Set column widths for component details
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(4.5)
        
        details = [
            ("Component Name", component_profile.get('name', '')),
            ("Component Type", component_profile.get('type', '')),
            ("Application", component_profile.get('application', '')),
            ("Test Level", component_profile.get('test_level', '')),
            ("Variants", ", ".join(component_profile.get('variants', [])))
        ]
        
        for key, value in details:
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
            
        # 2. Test Plan Matrix
        document.add_heading('2. Test Plan Matrix', level=1)
        document.add_paragraph("The following table details the required tests, procedures, and acceptance criteria.")
        
        # Create table for test cases
        headers = ['ID', 'Test Description', 'Test Procedure', 'Acceptance Criteria', 'Qty']
        
        table = document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Set specific column widths for the matrix
        table.columns[0].width = Inches(0.5) # ID
        table.columns[1].width = Inches(1.5) # Description
        table.columns[2].width = Inches(2.5) # Procedure
        table.columns[3].width = Inches(1.5) # Criteria
        table.columns[4].width = Inches(0.5) # Qty
        
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # Apply shading and bold to header
            tc_pr = hdr_cells[i]._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '366092') # Dark blue
            tc_pr.append(shd)
            
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255) # White text
                
        for idx, test_case in enumerate(test_cases, start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = f"B{idx}"
            row_cells[1].text = str(test_case.get('test_name', test_case.get('test_description', '')))
            
            # Map 'detailed_procedure' from LLM to 'Procedure' column
            procedure = test_case.get('detailed_procedure', test_case.get('test_procedure', ''))
            row_cells[2].text = str(procedure)
            
            row_cells[3].text = str(test_case.get('acceptance_criteria', ''))
            row_cells[4].text = str(test_case.get('quantity', '5'))
            
        # 3. Traceability
        if include_traceability:
            document.add_heading('3. Requirement Traceability', level=1)
            document.add_paragraph("This section maps test cases to source requirements.")
            
            headers = ['Test ID', 'Requirement ID', 'Source Standard']
            table = document.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                # Apply shading and bold to header
                tc_pr = hdr_cells[i]._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), '366092') # Dark blue
                tc_pr.append(shd)
                
                for run in hdr_cells[i].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255) # White text
                
            for idx, test_case in enumerate(test_cases, start=1):
                traceability = test_case.get('traceability', {})
                row_cells = table.add_row().cells
                row_cells[0].text = f"B{idx}"
                row_cells[1].text = str(traceability.get('requirement_id', ''))
                
                source = f"{traceability.get('source_standard', '')} {traceability.get('source_clause', '')}"
                row_cells[2].text = source.strip()
                
        # 4. References
        document.add_heading('4. References', level=1)
        
        references = set()
        for test_case in test_cases:
            traceability = test_case.get('traceability', {})
            source_std = traceability.get('source_standard', '')
            source_clause = traceability.get('source_clause', '')

            if source_std:
                references.add(source_std)
                
        for ref in sorted(references):
            document.add_paragraph(ref, style='List Bullet')
            
        # Save file
        output_filename = f"DVP_Doc_{component_profile.get('name', 'Component').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = Path(settings.output_dir) / output_filename
        
        document.save(str(output_path))
        logger.info(f"DVP Document saved to: {output_path}")
        
        return str(output_path)

    def _create_test_matrix_sheet(self, component_profile: Dict[str, Any],
                                  test_cases: List[Dict[str, Any]]):
        """
        Create main test matrix sheet (Annex B format)
        """
        ws = self.workbook.create_sheet("Annex B-Electronics DVP", 0)

        # Header section
        ws['A1'] = f"PROJECT NAME: {component_profile.get('name', 'Component')}"
        ws['A1'].font = Font(bold=True, size=14)

        ws['A2'] = f"Component Type: {component_profile.get('type', '')}"
        ws['A3'] = f"Application: {component_profile.get('application', '')}"
        ws['A4'] = f"Test Level: {component_profile.get('test_level', '')}"

        # Column headers (row 6)
        headers = [
            'Sl.No.',
            'Test Standard',
            'Test Description',
            'Test Procedure',
            'Acceptance Criteria',
            'Test Responsibility',
            'Test Stage',
            'Qty',
            'Test Days',
            'Start date',
            'End date',
            'Test Inference',
            'PCB/LAMP ASSEMBLY',
            'Remarks'
        ]

        for col_idx, header in enumerate(headers, start=1):
            cell = ws.cell(row=6, column=col_idx)
            cell.value = header
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        # Data rows
        for row_idx, test_case in enumerate(test_cases, start=7):
            ws.cell(row=row_idx, column=1).value = test_case.get('test_id', f'B{row_idx-6}')
            ws.cell(row=row_idx, column=2).value = test_case.get('test_standard', '')
            ws.cell(row=row_idx, column=3).value = test_case.get('test_description', '')
            ws.cell(row=row_idx, column=4).value = test_case.get('test_procedure', '')
            ws.cell(row=row_idx, column=5).value = test_case.get('acceptance_criteria', '')
            ws.cell(row=row_idx, column=6).value = test_case.get('test_responsibility', 'Supplier')
            ws.cell(row=row_idx, column=7).value = test_case.get('test_stage', 'DVP')
            ws.cell(row=row_idx, column=8).value = test_case.get('quantity', '')
            ws.cell(row=row_idx, column=9).value = test_case.get('estimated_days', 5)
            ws.cell(row=row_idx, column=10).value = ''  # Start date
            ws.cell(row=row_idx, column=11).value = ''  # End date
            ws.cell(row=row_idx, column=12).value = ''  # Test Inference
            ws.cell(row=row_idx, column=13).value = test_case.get('pcb_or_lamp', component_profile.get('test_level', ''))
            ws.cell(row=row_idx, column=14).value = test_case.get('remarks', '')

            # Wrap text for procedure and criteria
            ws.cell(row=row_idx, column=4).alignment = Alignment(wrap_text=True, vertical='top')
            ws.cell(row=row_idx, column=5).alignment = Alignment(wrap_text=True, vertical='top')

        # Adjust column widths
        column_widths = {
            'A': 10, 'B': 15, 'C': 25, 'D': 40, 'E': 40,
            'F': 15, 'G': 12, 'H': 20, 'I': 10, 'J': 12,
            'K': 12, 'L': 15, 'M': 25, 'N': 20
        }

        for col, width in column_widths.items():
            ws.column_dimensions[col].width = width

        # Set row heights
        for row_idx in range(7, 7 + len(test_cases)):
            ws.row_dimensions[row_idx].height = 60

    def _create_test_sequence_sheet(self, test_cases: List[Dict[str, Any]]):
        """
        Create test sequence sheet
        """
        ws = self.workbook.create_sheet("EMC & ENV TEST SEQUENCE", 1)

        ws['A1'] = "TEST SEQUENCE"
        ws['A1'].font = Font(bold=True, size=14)

        # Group tests by category
        thermal_tests = [tc for tc in test_cases if 'thermal' in tc.get('test_description', '').lower()]
        mechanical_tests = [tc for tc in test_cases if 'mechanical' in tc.get('test_description', '').lower()]
        env_tests = [tc for tc in test_cases if 'environment' in tc.get('test_description', '').lower() or 'humidity' in tc.get('test_description', '').lower()]

        row = 3
        ws.cell(row=row, column=1).value = "ENVIRONMENTAL TESTS"
        ws.cell(row=row, column=1).font = Font(bold=True)

        row += 1
        for idx, test in enumerate(env_tests + thermal_tests, start=1):
            ws.cell(row=row, column=1).value = f"Leg {idx}"
            ws.cell(row=row, column=2).value = test.get('test_id', '')
            ws.cell(row=row, column=3).value = test.get('test_description', '')
            row += 1

        row += 2
        ws.cell(row=row, column=1).value = "MECHANICAL TESTS"
        ws.cell(row=row, column=1).font = Font(bold=True)

        row += 1
        for idx, test in enumerate(mechanical_tests, start=1):
            ws.cell(row=row, column=1).value = f"Leg {idx}"
            ws.cell(row=row, column=2).value = test.get('test_id', '')
            ws.cell(row=row, column=3).value = test.get('test_description', '')
            row += 1

    def _create_traceability_sheet(self, test_cases: List[Dict[str, Any]]):
        """
        Create traceability matrix sheet
        """
        ws = self.workbook.create_sheet("Traceability Matrix", 2)

        # Headers
        headers = [
            'Test ID',
            'Test Description',
            'Requirement ID',
            'Source Clause',
            'Source Standard',
            'Requirement Type',
            'Confidence Score',
            'Linking Method'
        ]

        for col_idx, header in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col_idx)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")

        # Data rows
        row_idx = 2
        for test_case in test_cases:
            traceability = test_case.get('traceability', {})

            ws.cell(row=row_idx, column=1).value = test_case.get('test_id', '')
            ws.cell(row=row_idx, column=2).value = test_case.get('test_description', '')
            ws.cell(row=row_idx, column=3).value = traceability.get('requirement_id', '')
            ws.cell(row=row_idx, column=4).value = traceability.get('source_clause', '')
            ws.cell(row=row_idx, column=5).value = traceability.get('source_standard', '')
            ws.cell(row=row_idx, column=6).value = traceability.get('requirement_type', '')
            ws.cell(row=row_idx, column=7).value = traceability.get('confidence_score', '')
            ws.cell(row=row_idx, column=8).value = 'Hybrid (Semantic + Graph)'

            row_idx += 1

        # Adjust column widths
        for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']:
            ws.column_dimensions[col].width = 20

    def _create_references_sheet(self, test_cases: List[Dict[str, Any]]):
        """
        Create source references sheet
        """
        ws = self.workbook.create_sheet("Source References", 3)

        ws['A1'] = "SOURCE STANDARDS AND CLAUSES REFERENCED"
        ws['A1'].font = Font(bold=True, size=12)

        # Collect unique references
        references = set()
        for test_case in test_cases:
            traceability = test_case.get('traceability', {})
            source_std = traceability.get('source_standard', '')
            source_clause = traceability.get('source_clause', '')

            if source_std and source_clause:
                references.add(f"{source_std} - Clause {source_clause}")

        # List references
        row = 3
        for ref in sorted(references):
            ws.cell(row=row, column=1).value = ref
            row += 1

async def process_dvp_generation(job_id: str, request: DVPGenerationRequest):
    """
    Background task for DVP generation
    """
    try:
        dvp_jobs[job_id]['status'] = JobStatus.PROCESSING
        dvp_jobs[job_id]['current_step'] = 'Generating DVP document'
        dvp_jobs[job_id]['progress_percent'] = 20.0

        generator = DVPGenerator()
        
        output_format = request.output_format.lower()
        if output_format == 'docx' or output_format == 'doc':
             output_path = generator.generate_dvp_docx(
                component_profile=request.component_profile.model_dump(),
                test_cases=request.test_cases,
                include_traceability=request.include_traceability_sheet
            )
        else:
            # Default to xlsx
            output_path = generator.generate_dvp(
                component_profile=request.component_profile.model_dump(),
                test_cases=request.test_cases,
                include_traceability=request.include_traceability_sheet
            )

        # Get file size
        file_size = Path(output_path).stat().st_size

        # Create DVP ID
        dvp_id = Path(output_path).stem

        # Store DVP metadata
        generated_dvps[dvp_id] = {
            'dvp_id': dvp_id,
            'file_path': output_path,
            'component_name': request.component_profile.name,
            'test_cases_count': len(request.test_cases),
            'file_size_bytes': file_size,
            'created_at': datetime.utcnow()
        }

        # Update job status
        dvp_jobs[job_id]['status'] = JobStatus.COMPLETED
        dvp_jobs[job_id]['current_step'] = 'Completed'
        dvp_jobs[job_id]['progress_percent'] = 100.0
        dvp_jobs[job_id]['result'] = {
            'dvp_id': dvp_id,
            'file_path': output_path,
            'file_size_bytes': file_size,
            'test_cases_count': len(request.test_cases)
        }

        logger.info(f"DVP generation job {job_id} completed: {dvp_id}")

    except Exception as e:
        logger.exception(f"DVP generation job {job_id} failed: {e}")
        dvp_jobs[job_id]['status'] = JobStatus.FAILED
        dvp_jobs[job_id]['error'] = str(e)

# ==================== ENDPOINTS ====================

@router.post("/generate", response_model=DVPGenerationResponse)
async def generate_dvp_document(
    request: DVPGenerationRequest,
    background_tasks: BackgroundTasks
):
    """
    **Endpoint 5: Generate DVP Document**

    Creates complete DVP document (Excel or Word) matching reference format.

    **Sheets Generated (Excel):**
    1. Annex B - Electronics DVP (main test matrix)
    2. EMC & ENV TEST SEQUENCE (test grouping)
    3. Traceability Matrix (requirement → test mapping)
    4. Source References (standards cited)

    **Format:**
    - Excel .xlsx format
    - Word .docx format (Descriptive Test Plan)
    - Matches reference DVP layout
    - Includes styling and formatting
    - Ready for use by test engineers

    **Parameters:**
    - component_profile: Component specifications
    - test_cases: Generated test cases from LLM
    - output_format: xlsx (json, pdf future)
    - include_traceability_sheet: Include traceability

    **Example:**
    ```json
    {
        "component_profile": {...},
        "test_cases": [...],
        "output_format": "xlsx",
        "include_traceability_sheet": true
    }
    ```
    """
    job_id = str(uuid.uuid4())

    # Validate test cases
    if not request.test_cases:
        raise HTTPException(
            status_code=400,
            detail="No test cases provided. Please generate test cases first using /llm/generate"
        )

    # Create job entry
    dvp_jobs[job_id] = {
        'job_id': job_id,
        'status': JobStatus.PENDING,
        'current_step': 'Initializing',
        'progress_percent': 0.0,
        'created_at': datetime.utcnow()
    }

    # Start background processing
    background_tasks.add_task(
        process_dvp_generation,
        job_id,
        request
    )

    return DVPGenerationResponse(
        job_id=job_id,
        dvp_id="",  # Will be set when completed
        status=JobStatus.PENDING,
        message="DVP generation started. Use /dvp/status/{job_id} to check progress.",
        download_url="",
        file_size_bytes=0,
        test_cases_count=len(request.test_cases),
        requirements_covered=0,
        traceability_complete=request.include_traceability_sheet,
        timestamp=datetime.utcnow()
    )

@router.get("/status/{job_id}", response_model=JobStatusResponse)
async def get_dvp_generation_status(job_id: str):
    """
    **Check DVP generation job status**

    **Parameters:**
    - job_id: Job ID from /generate endpoint

    **Returns:**
    - Current status and progress
    - DVP ID and download URL when completed
    """
    if job_id not in dvp_jobs:
        raise HTTPException(status_code=404, detail=f"Job {job_id} not found")

    job = dvp_jobs[job_id]

    return JobStatusResponse(
        job_id=job_id,
        status=job['status'],
        progress_percent=job.get('progress_percent', 0.0),
        current_step=job.get('current_step', 'Unknown'),
        message=f"DVP generation: {job.get('current_step', 'Processing')}",
        result=job.get('result') if job['status'] == JobStatus.COMPLETED else None,
        error=job.get('error')
    )

@router.get("/download/{dvp_id}")
async def download_dvp(dvp_id: str):
    """
    **Endpoint 6: Download Generated DVP**

    Download the Excel DVP document.

    **Parameters:**
    - dvp_id: DVP ID from generation response

    **Returns:**
    - Excel file download
    """
    if dvp_id not in generated_dvps:
        raise HTTPException(status_code=404, detail=f"DVP {dvp_id} not found")

    dvp_metadata = generated_dvps[dvp_id]
    file_path = dvp_metadata['file_path']

    if not Path(file_path).exists():
        raise HTTPException(status_code=404, detail="DVP file not found on disk")

    return FileResponse(
        path=file_path,
        filename=Path(file_path).name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if str(file_path).endswith('.docx') else "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

@router.get("/list")
async def list_generated_dvps():
    """
    **List all generated DVPs**

    Returns a list of all DVP documents with metadata.
    """
    return {
        "total_dvps": len(generated_dvps),
        "dvps": [
            {
                "dvp_id": dvp_id,
                "component_name": metadata['component_name'],
                "test_cases_count": metadata['test_cases_count'],
                "file_size_bytes": metadata['file_size_bytes'],
                "created_at": metadata['created_at'].isoformat(),
                "download_url": f"/api/v1/dvp/download/{dvp_id}"
            }
            for dvp_id, metadata in generated_dvps.items()
        ]
    }

@router.delete("/delete/{dvp_id}")
async def delete_dvp(dvp_id: str):
    """
    **Delete a generated DVP**

    Removes the DVP file and metadata.

    **Parameters:**
    - dvp_id: DVP ID to delete
    """
    if dvp_id not in generated_dvps:
        raise HTTPException(status_code=404, detail=f"DVP {dvp_id} not found")

    dvp_metadata = generated_dvps[dvp_id]
    file_path = Path(dvp_metadata['file_path'])

    # Delete file
    if file_path.exists():
        file_path.unlink()

    # Remove from metadata
    del generated_dvps[dvp_id]

    return {
        "message": f"DVP {dvp_id} deleted successfully",
        "dvp_id": dvp_id
    }
