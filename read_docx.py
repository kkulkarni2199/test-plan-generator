from docx import Document
import sys

def read_docx(file_path):
    try:
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
        doc = Document(file_path)
        print(f"File: {file_path}")
        print("="*40)
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"[{i}] {para.text}")
        
        print("\nTABLES:")
        for t_idx, table in enumerate(doc.tables):
            print(f"\nTable {t_idx+1}:")
            for r_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                print(f"Row {r_idx+1}: {' | '.join(cells)}")
        print("="*40)
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        read_docx(sys.argv[1])
    else:
        print("Usage: python read_docx.py <path_to_docx>")
